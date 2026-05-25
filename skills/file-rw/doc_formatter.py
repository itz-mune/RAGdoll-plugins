"""
doc_formatter.py — converts Markdown content into richly formatted documents.

Dispatch via:  render_document(ext, markdown_content, output_path)

Supported output formats
------------------------
docx  — python-docx  (headings, paragraphs, lists, tables, code blocks,
                       bold / italic / inline-code, blockquotes, HR,
                       accent-coloured heading styles, styled table headers)
pdf   — WeasyPrint   (primary: full CSS, A4 @page rules, full-bleed blue
                       title block, page numbers via CSS counters, Unicode)
        reportlab    (fallback: BaseDocTemplate, drawn title block,
                       running header/footer, Unicode-sanitised text)
html  — standalone   (embedded CSS, responsive, automatic dark-mode)

All renderers share the same block-level Markdown parser so the output
is visually consistent across formats.
"""
from __future__ import annotations

import re
from pathlib import Path

# Shared accent colour (professional blue) used across every format
_ACCENT_HEX = "#2E74B5"
_ACCENT_RGB = (0x2E, 0x74, 0xB5)


# ══════════════════════════════════════════════════════════════════════════════
# Shared: block-level Markdown parser
# ══════════════════════════════════════════════════════════════════════════════

def _parse_blocks(content: str) -> list[dict]:
    """
    Parse Markdown into a flat list of typed block dicts.

    Block types
    -----------
    heading   : {type, level: int, text: str}
    paragraph : {type, text: str}
    bullet    : {type, text: str, depth: int}
    numbered  : {type, text: str, depth: int}
    quote     : {type, text: str}
    code      : {type, lang: str, lines: list[str]}
    hr        : {type}
    table     : {type, rows: list[list[str]]}  — separator rows excluded
    blank     : {type}
    """
    blocks: list[dict] = []
    lines = content.splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i]

        # blank
        if not raw.strip():
            blocks.append({"type": "blank"})
            i += 1
            continue

        # fenced code block
        if raw.lstrip().startswith("```"):
            lang = raw.strip().lstrip("`").strip()
            body: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                body.append(lines[i])
                i += 1
            blocks.append({"type": "code", "lang": lang, "lines": body})
            i += 1  # skip closing ```
            continue

        # horizontal rule
        if re.match(r"^\s*[-*_]{3,}\s*$", raw):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # ATX heading
        m = re.match(r"^(#{1,6})\s+(.*)", raw)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue

        # table
        if "|" in raw and raw.strip().startswith("|"):
            table_raw: list[str] = []
            while i < len(lines) and "|" in lines[i]:
                table_raw.append(lines[i])
                i += 1
            rows = []
            for row_line in table_raw:
                if re.match(r"^\|[\s\-:|]+\|$", row_line.strip()):
                    continue  # separator row
                cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
                rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        # bullet list
        m = re.match(r"^(\s*)[-*+]\s+(.*)", raw)
        if m:
            depth = len(m.group(1)) // 2
            blocks.append({"type": "bullet", "text": m.group(2), "depth": depth})
            i += 1
            continue

        # numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", raw)
        if m:
            depth = len(m.group(1)) // 2
            blocks.append({"type": "numbered", "text": m.group(2), "depth": depth})
            i += 1
            continue

        # blockquote
        m = re.match(r"^>\s?(.*)", raw)
        if m:
            blocks.append({"type": "quote", "text": m.group(1)})
            i += 1
            continue

        # image  ![alt](src)  or  ![alt|align](src)
        # align hint: left | right | center | full  (omit for auto-detect)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", raw.strip())
        if m:
            alt_raw = m.group(1)
            src     = m.group(2).strip()
            _ALIGNS = {"left", "right", "center", "full"}
            if "|" in alt_raw:
                _parts = alt_raw.rsplit("|", 1)
                _hint  = _parts[1].strip().lower()
                if _hint in _ALIGNS:
                    alt, align = _parts[0].strip(), _hint
                else:
                    alt, align = alt_raw, "auto"
            else:
                alt, align = alt_raw, "auto"
            blocks.append({"type": "image", "alt": alt, "src": src, "align": align})
            i += 1
            continue

        # paragraph
        blocks.append({"type": "paragraph", "text": raw.rstrip()})
        i += 1

    return blocks


# ══════════════════════════════════════════════════════════════════════════════
# Shared: inline Markdown → spans  (for docx)
# ══════════════════════════════════════════════════════════════════════════════

_INLINE_RE = re.compile(
    r"(\*\*\*(?P<biu>[^*]+)\*\*\*"
    r"|\*\*(?P<bi>[^*]+)\*\*"
    r"|\*(?P<i>[^*]+)\*"
    r"|__(?P<b2>[^_]+)__"
    r"|_(?P<i2>[^_]+)_"
    r"|`(?P<code>[^`]+)`)"
)

def _inline_spans(text: str) -> list[dict]:
    """Return a list of {text, bold, italic, code} dicts."""
    spans: list[dict] = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            spans.append({"text": text[pos:m.start()], "bold": False, "italic": False, "code": False})
        if m.group("biu"):
            spans.append({"text": m.group("biu"),          "bold": True,  "italic": True,  "code": False})
        elif m.group("bi"):
            spans.append({"text": m.group("bi"),           "bold": True,  "italic": False, "code": False})
        elif m.group("i") or m.group("i2"):
            spans.append({"text": m.group("i") or m.group("i2"), "bold": False, "italic": True,  "code": False})
        elif m.group("b2"):
            spans.append({"text": m.group("b2"),           "bold": True,  "italic": False, "code": False})
        elif m.group("code"):
            spans.append({"text": m.group("code"),         "bold": False, "italic": False, "code": True})
        pos = m.end()
    if pos < len(text):
        spans.append({"text": text[pos:], "bold": False, "italic": False, "code": False})
    return spans or [{"text": text, "bold": False, "italic": False, "code": False}]


# ══════════════════════════════════════════════════════════════════════════════
# DOCX renderer  (python-docx)
# ══════════════════════════════════════════════════════════════════════════════

def markdown_to_docx(content: str, path: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed — run: uv add python-docx")

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # Base styles
    accent = RGBColor(*_ACCENT_RGB)
    _docx_set_style(doc.styles["Normal"], "Calibri", Pt(11))
    for level, size in ((1, Pt(20)), (2, Pt(16)), (3, Pt(14)), (4, Pt(13))):
        sname = f"Heading {level}"
        if sname in doc.styles:
            _docx_set_style(doc.styles[sname], "Calibri", size, bold=True, colour=accent)

    # Render blocks
    for blk in _parse_blocks(content):
        t = blk["type"]

        if t == "blank":
            continue

        elif t == "heading":
            lvl  = min(blk["level"], 4)
            para = doc.add_heading("", level=lvl)
            para.clear()
            _docx_inline_runs(para, blk["text"])

        elif t == "paragraph":
            para = doc.add_paragraph()
            _docx_inline_runs(para, blk["text"])

        elif t == "quote":
            para = doc.add_paragraph()
            _docx_inline_runs(para, blk["text"])
            para.paragraph_format.left_indent = Inches(0.4)
            for run in para.runs:
                run.italic = True

        elif t == "bullet":
            style = "List Bullet 2" if blk.get("depth", 0) else "List Bullet"
            para  = doc.add_paragraph(style=style)
            _docx_inline_runs(para, blk["text"])

        elif t == "numbered":
            style = "List Number 2" if blk.get("depth", 0) else "List Number"
            para  = doc.add_paragraph(style=style)
            _docx_inline_runs(para, blk["text"])

        elif t == "code":
            code_text = "\n".join(blk["lines"])
            para = doc.add_paragraph()
            run  = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            # grey background shading
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F2F2F2")
            pPr.append(shd)
            para.paragraph_format.left_indent  = Inches(0.3)
            para.paragraph_format.right_indent = Inches(0.3)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)

        elif t == "hr":
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "AAAAAA")
            pBdr.append(bot)
            pPr.append(pBdr)

        elif t == "image":
            img_bytes = _fetch_image_bytes(blk["src"])
            if img_bytes:
                try:
                    align = _resolve_align(blk.get("align", "auto"), img_bytes)
                    _docx_add_image(doc, img_bytes, align)
                    if blk.get("alt"):
                        cap = doc.add_paragraph(blk["alt"])
                        cap.paragraph_format.alignment = 1
                        for _run in cap.runs:
                            _run.italic = True
                            _run.font.size = Pt(9)
                except Exception:
                    pass  # skip silently if image can't be embedded

        elif t == "table":
            rows = blk["rows"]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl   = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell_text = row[ci] if ci < len(row) else ""
                    cell      = tbl.cell(ri, ci)
                    cell.text = ""
                    para      = cell.paragraphs[0]
                    _docx_inline_runs(para, cell_text)
                    if ri == 0:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd  = OxmlElement("w:shd")
                        shd.set(qn("w:val"),   "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"),  "2E74B5")
                        tcPr.append(shd)

    doc.save(path)


def _docx_add_image(doc, img_bytes: bytes, align: str, text_w_inches: float = 6.1) -> None:
    """Add an image to a DOCX document with the given alignment.

    align == 'full'   → full text-area width, inline
    align == 'center' → 68 % width, centred paragraph
    align == 'left'   → 42 % width, float-left via wp:anchor (text wraps right)
    align == 'right'  → 42 % width, float-right via wp:anchor (text wraps left)
    """
    from docx.shared import Inches, Pt
    from io import BytesIO

    width = _smart_width_inches(align, text_w_inches)
    para  = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)

    if align == "center":
        para.paragraph_format.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.add_picture(BytesIO(img_bytes), width=Inches(width))

    if align not in ("left", "right"):
        return  # inline rendering is correct for full/center

    # ── Convert wp:inline → wp:anchor so Word wraps text around the image ──
    try:
        import copy
        import lxml.etree as etree
        from docx.oxml.ns import qn

        WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        A  = "http://schemas.openxmlformats.org/drawingml/2006/main"

        drawing_el = run._r.find(qn("w:drawing"))
        if drawing_el is None:
            return
        inline = drawing_el.find(f"{{{WP}}}inline")
        if inline is None:
            return

        # Harvest key child elements from the inline image
        extent  = inline.find(f"{{{WP}}}extent")
        doc_pr  = inline.find(f"{{{WP}}}docPr")
        cnv_pr  = inline.find(f"{{{WP}}}cNvGraphicFramePr")
        graphic = inline.find(f"{{{A}}}graphic")

        cx = extent.get("cx") if extent is not None else str(int(width * 914400))
        cy = extent.get("cy") if extent is not None else str(int(width * 914400 * 0.75))
        pr_id   = (doc_pr.get("id",   "1")       if doc_pr is not None else "1")
        pr_name = (doc_pr.get("name", "Image")   if doc_pr is not None else "Image")

        # Build wp:anchor in the correct OOXML child order
        anc = etree.Element(f"{{{WP}}}anchor")
        for k, v in [
            ("distT", "0"), ("distB", "0"),
            ("distL", "114300"), ("distR", "114300"),
            ("simplePos", "0"), ("relativeHeight", "251658240"),
            ("behindDoc", "0"), ("locked", "0"),
            ("layoutInCell", "1"), ("allowOverlap", "0"),
        ]:
            anc.set(k, v)

        sp = etree.SubElement(anc, f"{{{WP}}}simplePos")
        sp.set("x", "0"); sp.set("y", "0")

        ph = etree.SubElement(anc, f"{{{WP}}}positionH")
        ph.set("relativeFrom", "column")
        etree.SubElement(ph, f"{{{WP}}}align").text = align   # "left" or "right"

        pv = etree.SubElement(anc, f"{{{WP}}}positionV")
        pv.set("relativeFrom", "paragraph")
        etree.SubElement(pv, f"{{{WP}}}align").text = "top"

        ext = etree.SubElement(anc, f"{{{WP}}}extent")
        ext.set("cx", cx); ext.set("cy", cy)

        eff = etree.SubElement(anc, f"{{{WP}}}effectExtent")
        for s in ("l", "t", "r", "b"):
            eff.set(s, "0")

        wrap = etree.SubElement(anc, f"{{{WP}}}wrapSquare")
        wrap.set("wrapText", "bothSides")

        dp = etree.SubElement(anc, f"{{{WP}}}docPr")
        dp.set("id", pr_id); dp.set("name", pr_name)

        if cnv_pr is not None:
            anc.append(copy.deepcopy(cnv_pr))

        if graphic is not None:
            anc.append(copy.deepcopy(graphic))

        drawing_el.remove(inline)
        drawing_el.append(anc)

    except Exception:
        pass  # inline image already present — safe fallback


def _docx_set_style(style, font_name: str, size, bold: bool = False, colour=None) -> None:
    style.font.name = font_name
    style.font.size = size
    if bold:
        style.font.bold = True
    if colour is not None:
        style.font.color.rgb = colour


def _docx_inline_runs(para, text: str) -> None:
    """Add inline-formatted runs to a python-docx paragraph."""
    from docx.shared import Pt
    for span in _inline_spans(text):
        run = para.add_run(span["text"])
        if span["bold"]:
            run.bold = True
        if span["italic"]:
            run.italic = True
        if span["code"]:
            run.font.name = "Courier New"
            run.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════════════════════
# Shared inline-HTML helper  (used by both HTML and PDF renderers)
# ══════════════════════════════════════════════════════════════════════════════

def _inline_html_safe(text: str) -> str:
    t = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    t = re.sub(r"\*\*\*([^*]+)\*\*\*", r"<strong><em>\1</em></strong>", t)
    t = re.sub(r"\*\*([^*]+)\*\*",     r"<strong>\1</strong>",           t)
    t = re.sub(r"__([^_]+)__",         r"<strong>\1</strong>",           t)
    t = re.sub(r"\*([^*]+)\*",         r"<em>\1</em>",                   t)
    t = re.sub(r"_([^_]+)_",           r"<em>\1</em>",                   t)
    t = re.sub(r"`([^`]+)`",           r"<code>\1</code>",               t)
    t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>',       t)
    return t


def _fetch_image_bytes(src: str) -> bytes | None:
    """Download an image from a URL or read it from a local path.

    Returns bytes normalised to JPEG or PNG via Pillow (if installed),
    or the raw bytes as-is if Pillow is unavailable. Returns None on any error.
    """
    try:
        if src.startswith(("http://", "https://")):
            import urllib.request
            req = urllib.request.Request(src, headers={"User-Agent": "RAGdoll/1.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read()
        else:
            raw = Path(src).read_bytes()

        # Normalise via Pillow so any source format works everywhere
        try:
            from PIL import Image as _PILImage
            from io import BytesIO as _BytesIO
            img = _PILImage.open(_BytesIO(raw))
            out = _BytesIO()
            if img.mode == "RGBA":
                img.save(out, format="PNG")
            else:
                img = img.convert("RGB")
                img.save(out, format="JPEG", quality=85)
            return out.getvalue()
        except ImportError:
            return raw  # Pillow not available — pass through as-is
    except Exception:
        return None


def _aspect_ratio(img_bytes: bytes) -> float | None:
    """Return width/height aspect ratio using Pillow, or None if unavailable."""
    try:
        from PIL import Image as _PIL
        from io import BytesIO as _BIO
        w, h = _PIL.open(_BIO(img_bytes)).size
        return w / h if h else None
    except Exception:
        return None


def _resolve_align(align_hint: str, img_bytes: bytes | None) -> str:
    """Turn an 'auto' hint (or any explicit one) into a concrete alignment word.

    Rules (applied only when hint == 'auto'):
      ratio >= 2.0   →  full   (panoramic/banner)
      ratio >= 1.05  →  center (normal landscape)
      ratio <  1.05  →  right  (portrait/square — floats right, text wraps left)
    """
    if align_hint != "auto":
        return align_hint
    if img_bytes is None:
        return "center"
    ratio = _aspect_ratio(img_bytes)
    if ratio is None:
        return "center"
    if ratio >= 2.0:
        return "full"
    if ratio >= 1.05:
        return "center"
    return "right"


def _smart_width_inches(align: str, text_w: float = 5.5) -> float:
    """Return display width in inches for the given alignment and text area."""
    if align == "full":
        return text_w
    if align == "center":
        return round(text_w * 0.68, 2)   # ~68 % — nicely centred
    # left / right float
    return round(text_w * 0.42, 2)       # ~42 % — leaves room for text


def _html_fig_class(align: str) -> str:
    return {
        "full":   "img-full",
        "center": "img-center",
        "left":   "img-wrap-left",
        "right":  "img-wrap-right",
    }.get(align, "img-center")


def _blocks_to_html_str(blocks: list[dict]) -> str:
    """Render a list of parsed blocks to an HTML fragment string."""
    parts: list[str] = []
    in_ul = in_ol = False

    for blk in blocks:
        btype = blk["type"]
        if btype == "bullet":
            if not in_ul: parts.append("<ul>"); in_ul = True
            if in_ol:     parts.append("</ol>"); in_ol = False
        elif btype == "numbered":
            if not in_ol: parts.append("<ol>"); in_ol = True
            if in_ul:     parts.append("</ul>"); in_ul = False
        else:
            if in_ul: parts.append("</ul>"); in_ul = False
            if in_ol: parts.append("</ol>"); in_ol = False

        if   btype == "blank":   pass
        elif btype == "heading":
            lvl = min(blk["level"], 6)
            parts.append(f"<h{lvl}>{_inline_html_safe(blk['text'])}</h{lvl}>")
        elif btype == "paragraph":
            parts.append(f"<p>{_inline_html_safe(blk['text'])}</p>")
        elif btype == "quote":
            parts.append(f"<blockquote><p>{_inline_html_safe(blk['text'])}</p></blockquote>")
        elif btype in ("bullet", "numbered"):
            parts.append(f"<li>{_inline_html_safe(blk['text'])}</li>")
        elif btype == "code":
            lang = blk.get("lang", "")
            cls  = f' class="language-{lang}"' if lang else ""
            body = "\n".join(blk["lines"]).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            parts.append(f"<pre><code{cls}>{body}</code></pre>")
        elif btype == "hr":
            parts.append("<hr>")
        elif btype == "image":
            src        = blk["src"]
            alt        = blk.get("alt", "")
            align_hint = blk.get("align", "auto")
            alt_esc    = alt.replace('"', "&quot;")
            caption    = f"<figcaption>{_inline_html_safe(alt)}</figcaption>" if alt else ""

            if src.startswith(("http://", "https://")):
                # Fetch to determine aspect ratio for auto-align; use URL directly in HTML
                img_bytes_hint = _fetch_image_bytes(src) if align_hint == "auto" else None
                align    = _resolve_align(align_hint, img_bytes_hint)
                fig_cls  = _html_fig_class(align)
                parts.append(
                    f'<figure class="{fig_cls}"><img src="{src}" alt="{alt_esc}">'
                    f'{caption}</figure>'
                )
            else:
                img_bytes = _fetch_image_bytes(src)
                if img_bytes:
                    import base64
                    align   = _resolve_align(align_hint, img_bytes)
                    fig_cls = _html_fig_class(align)
                    mime    = "image/png" if img_bytes[:4] == b"\x89PNG" else "image/jpeg"
                    b64     = base64.b64encode(img_bytes).decode()
                    parts.append(
                        f'<figure class="{fig_cls}"><img src="data:{mime};base64,{b64}" alt="{alt_esc}">'
                        f'{caption}</figure>'
                    )
        elif btype == "table":
            rows  = blk["rows"]
            ncols = max(len(r) for r in rows)
            tbl   = ["<table>"]
            for ri, row in enumerate(rows):
                tbl.append("<tr>")
                for ci in range(ncols):
                    cell = row[ci] if ci < len(row) else ""
                    tag  = "th" if ri == 0 else "td"
                    tbl.append(f"<{tag}>{_inline_html_safe(cell)}</{tag}>")
                tbl.append("</tr>")
            tbl.append("</table>")
            parts.append("\n".join(tbl))

    if in_ul: parts.append("</ul>")
    if in_ol: parts.append("</ol>")
    return "\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
# PDF renderer  — WeasyPrint (primary) / reportlab (fallback)
# ══════════════════════════════════════════════════════════════════════════════

def markdown_to_pdf(content: str, path: str) -> None:
    """Render Markdown → PDF.
    Primary:  WeasyPrint  (HTML→PDF, full CSS, Unicode, page numbers, title block).
    Fallback: reportlab   (pure-Python, no system deps).
    """
    try:
        _pdf_weasyprint(content, path)
    except ImportError:
        _pdf_reportlab(content, path)
    except Exception:
        # WeasyPrint installed but failed at runtime — fall back gracefully
        try:
            _pdf_reportlab(content, path)
        except Exception as exc:
            raise exc


# ── WeasyPrint path ────────────────────────────────────────────────────────────

def _pdf_weasyprint(content: str, path: str) -> None:
    try:
        import weasyprint  # type: ignore
    except ImportError:
        raise ImportError("weasyprint")

    html = _build_pdf_html(content, path)
    weasyprint.HTML(string=html).write_pdf(path)


def _build_pdf_html(content: str, path: str) -> str:
    """Build a print-optimised, self-contained HTML string for WeasyPrint."""
    blocks = _parse_blocks(content)

    # Extract H1 as document title; body starts after it
    doc_title = Path(path).stem.replace("_", " ").replace("-", " ")
    body_blocks = blocks
    if blocks and blocks[0]["type"] == "heading" and blocks[0]["level"] == 1:
        doc_title = blocks[0]["text"]
        body_blocks = blocks[1:]

    title_html = _inline_html_safe(doc_title)
    body_html  = _blocks_to_html_str(body_blocks)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{title_html}</title>
<style>
{_PDF_CSS}
</style>
</head>
<body>
<div class="title-block">
  <div class="title-rule"></div>
  <p class="doc-title">{title_html}</p>
</div>
<article>
{body_html}
</article>
</body>
</html>"""


_PDF_CSS = """
/* ── Page layout ──────────────────────────────────────────────────────────── */
@page {
  size: A4;
  margin: 2.2cm 2.2cm 2.6cm;

  @bottom-center {
    content: counter(page);
    font-size: 9pt;
    color: #aaa;
    font-family: "Segoe UI", Helvetica, Arial, sans-serif;
  }
}
@page :first {
  margin-top: 0;
  @bottom-center { content: none; }
}

/* ── Base ─────────────────────────────────────────────────────────────────── */
*, *::before, *::after { box-sizing: border-box; }

body {
  font-family: "Segoe UI", "Calibri", "Helvetica Neue", Helvetica, Arial, sans-serif;
  font-size: 10.5pt;
  line-height: 1.65;
  color: #1a1a1a;
  background: white;
  margin: 0;
}

/* ── Title block (first page, bleeds to edges) ────────────────────────────── */
.title-block {
  background: #2E74B5;
  padding: 2.8cm 2.2cm 2cm;
  margin: 0 -2.2cm 1.8cm;
  page-break-after: avoid;
}

.title-rule {
  height: 3px;
  background: rgba(255,255,255,0.30);
  border-radius: 2px;
  margin-bottom: 1.1cm;
}

.doc-title {
  font-size: 22pt;
  font-weight: 700;
  color: white;
  margin: 0;
  line-height: 1.25;
  letter-spacing: -0.2pt;
}

/* ── Headings ─────────────────────────────────────────────────────────────── */
h1, h2, h3, h4, h5, h6 {
  color: #2E74B5;
  font-weight: 700;
  line-height: 1.3;
  margin-top: 1.6em;
  margin-bottom: 0.4em;
  page-break-after: avoid;
}
h1 { font-size: 16pt; border-bottom: 2px solid #2E74B5;  padding-bottom: 0.2em; }
h2 { font-size: 13pt; border-bottom: 1px solid #c8d8ea; padding-bottom: 0.15em; }
h3 { font-size: 11.5pt; color: #1A4F8A; }
h4 { font-size: 10.5pt; color: #1A4F8A; font-style: italic; }

/* ── Paragraphs ───────────────────────────────────────────────────────────── */
p { margin: 0.6em 0; text-align: justify; orphans: 2; widows: 2; }

a      { color: #2E74B5; }
strong { font-weight: 700; }
em     { font-style: italic; }

/* ── Inline code ──────────────────────────────────────────────────────────── */
code {
  font-family: "Cascadia Code", "Fira Code", Consolas, "Courier New", monospace;
  font-size: 0.85em;
  background: #eef2f7;
  border: 1px solid #c8d8ea;
  border-radius: 3px;
  padding: 0.05em 0.35em;
}

/* ── Code blocks ──────────────────────────────────────────────────────────── */
pre {
  background: #f5f7fa;
  border: 1px solid #c8d8ea;
  border-left: 4px solid #2E74B5;
  border-radius: 4px;
  padding: 0.85em 1em;
  margin: 1em 0;
  page-break-inside: avoid;
}
pre code {
  background: none;
  border: none;
  padding: 0;
  font-size: 8.5pt;
  line-height: 1.5;
}

/* ── Blockquote ───────────────────────────────────────────────────────────── */
blockquote {
  border-left: 4px solid #2E74B5;
  background: #f5f8fd;
  margin: 1em 0;
  padding: 0.6em 1em;
  color: #444;
  font-style: italic;
  page-break-inside: avoid;
}
blockquote p { margin: 0; text-align: left; }

/* ── Lists ────────────────────────────────────────────────────────────────── */
ul, ol { padding-left: 1.5em; margin: 0.5em 0; }
li     { margin: 0.25em 0; }

/* ── HR ───────────────────────────────────────────────────────────────────── */
hr { border: none; border-top: 1.5px solid #c8d8ea; margin: 1.5em 0; }

/* ── Tables ───────────────────────────────────────────────────────────────── */
table {
  width: 100%;
  border-collapse: collapse;
  margin: 1.2em 0;
  font-size: 9.5pt;
  page-break-inside: avoid;
}
th, td {
  border: 1px solid #c8d8ea;
  padding: 0.45em 0.75em;
  text-align: left;
  vertical-align: top;
}
th {
  background: #2E74B5;
  color: white;
  font-weight: 600;
  font-size: 9pt;
}
tr:nth-child(even) td { background: #f5f8fd; }

/* ── Images ───────────────────────────────────────────────────────────────── */
figure { page-break-inside: avoid; margin: 0; }
figure img { height: auto; border-radius: 4px; display: block; }
figcaption { font-size: 8pt; color: #777; margin-top: 0.3em; font-style: italic; text-align: center; }

/* Full-width */
figure.img-full  { width: 100%; margin: 1.4em 0; }
figure.img-full img { max-width: 100%; margin: 0 auto; }

/* Centred, partial width */
figure.img-center { width: 68%; margin: 1.4em auto; }
figure.img-center img { width: 100%; }

/* Float left — text wraps on the right */
figure.img-wrap-left {
  float: left;
  width: 42%;
  margin: 0.4em 1.5em 1em 0;
  clear: left;
}
figure.img-wrap-left img { width: 100%; }

/* Float right — text wraps on the left */
figure.img-wrap-right {
  float: right;
  width: 42%;
  margin: 0.4em 0 1em 1.5em;
  clear: right;
}
figure.img-wrap-right img { width: 100%; }

/* Headings and HR clear floats so the next section starts cleanly */
h1, h2, h3, h4, h5, h6, hr { clear: both; }
"""


# ── reportlab fallback ─────────────────────────────────────────────────────────

# Common Unicode characters that fall outside reportlab's default Latin-1 encoding
_RL_UNICODE_SUBS: dict[str, str] = {
    "—": "--",    # em dash
    "–": "-",     # en dash
    "‘": "'",     # left single quote
    "’": "'",     # right single quote
    "“": '"',     # left double quote
    "”": '"',     # right double quote
    "…": "...",   # ellipsis
    " ": " ",     # non-breaking space
    "•": "-",     # bullet (we add our own)
    "′": "'",     # prime
    "″": '"',     # double prime
    "°": " deg",  # degree sign
    "→": "->",    # right arrow
    "←": "<-",    # left arrow
    "×": "x",     # multiplication sign
    "÷": "/",     # division sign
    "≈": "~=",    # almost equal
    "≠": "!=",    # not equal
    "≤": "<=",    # less-than or equal
    "≥": ">=",    # greater-than or equal
}


def _rl_sanitize(text: str) -> str:
    """Replace common Unicode chars not supported by reportlab's built-in fonts."""
    for ch, rep in _RL_UNICODE_SUBS.items():
        text = text.replace(ch, rep)
    # Drop any remaining non-Latin-1 characters silently
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def _escape_pdf(text: str) -> str:
    return _rl_sanitize(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_pdf(text: str) -> str:
    """Convert inline Markdown to reportlab XML markup (Latin-1 safe)."""
    t = _escape_pdf(text)
    t = re.sub(r"\*\*\*([^*]+)\*\*\*", r"<b><i>\1</i></b>", t)
    t = re.sub(r"\*\*([^*]+)\*\*",     r"<b>\1</b>",         t)
    t = re.sub(r"__([^_]+)__",         r"<b>\1</b>",         t)
    t = re.sub(r"\*([^*]+)\*",         r"<i>\1</i>",         t)
    t = re.sub(r"_([^_]+)_",           r"<i>\1</i>",         t)
    t = re.sub(r"`([^`]+)`",           r'<font name="Courier">\1</font>', t)
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)
    return t


def _pdf_reportlab(content: str, path: str) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch, mm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER
        from reportlab.platypus import (
            BaseDocTemplate, PageTemplate, Frame,
            Paragraph, Spacer, Preformatted,
            Table, TableStyle, HRFlowable,
        )
    except ImportError:
        raise ImportError("reportlab not installed — run: uv add reportlab")

    # ── Extract title ──────────────────────────────────────────────────────────
    blocks = _parse_blocks(_rl_sanitize(content))
    doc_title = Path(path).stem.replace("_", " ").replace("-", " ")
    body_blocks = blocks
    if blocks and blocks[0]["type"] == "heading" and blocks[0]["level"] == 1:
        doc_title = _rl_sanitize(blocks[0]["text"])
        body_blocks = blocks[1:]

    # ── Colours ────────────────────────────────────────────────────────────────
    accent   = colors.HexColor(_ACCENT_HEX)
    grey_bg  = colors.HexColor("#F2F2F2")
    grey_alt = colors.HexColor("#F8FAFD")
    W, H     = A4
    ML = MR  = 1.2 * inch
    MT = MB  = 0.9 * inch
    TITLE_H  = 1.9 * inch   # height reserved for drawn title block on p.1

    # ── Page callbacks ─────────────────────────────────────────────────────────
    def _draw_title(canvas, _doc):
        canvas.saveState()
        # Blue title band
        canvas.setFillColor(accent)
        canvas.rect(0, H - TITLE_H, W, TITLE_H, fill=1, stroke=0)
        # Thin accent rule
        canvas.setFillColor(colors.HexColor("#FFFFFF"))
        canvas.setFillAlpha(0.3)
        canvas.rect(ML, H - 18*mm, W - ML - MR, 2.5, fill=1, stroke=0)
        canvas.setFillAlpha(1)
        # Title text
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 20)
        _rl_draw_wrapped(canvas, doc_title, ML, H - TITLE_H + 14*mm,
                         W - ML - MR, 20, "Helvetica-Bold", leading=26)
        _draw_footer(canvas, _doc.page)
        canvas.restoreState()

    def _draw_header_footer(canvas, _doc):
        canvas.saveState()
        # Running header line
        canvas.setStrokeColor(colors.HexColor("#D0DCE8"))
        canvas.setLineWidth(0.5)
        canvas.line(ML, H - MT + 3*mm, W - MR, H - MT + 3*mm)
        canvas.setFillColor(colors.HexColor("#888888"))
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(ML, H - MT + 5*mm, doc_title)
        _draw_footer(canvas, _doc.page)
        canvas.restoreState()

    def _draw_footer(canvas, page_num: int):
        canvas.setStrokeColor(colors.HexColor("#D0DCE8"))
        canvas.setLineWidth(0.5)
        canvas.line(ML, MB - 4*mm, W - MR, MB - 4*mm)
        canvas.setFillColor(colors.HexColor("#AAAAAA"))
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(W / 2, MB - 8*mm, str(page_num))

    def _rl_draw_wrapped(canvas, text, x, y, max_w, font_size, font_name, leading):
        """Naive word-wrap for canvas text drawing."""
        words = text.split()
        line  = ""
        lines_drawn = []
        for word in words:
            test = (line + " " + word).strip()
            if canvas.stringWidth(test, font_name, font_size) <= max_w:
                line = test
            else:
                lines_drawn.append(line)
                line = word
        if line:
            lines_drawn.append(line)
        for i, ln in enumerate(lines_drawn):
            canvas.drawString(x, y - i * leading, ln)

    # ── Page templates ─────────────────────────────────────────────────────────
    first_frame = Frame(ML, MB, W - ML - MR, H - TITLE_H - MB - 4*mm, id="first")
    later_frame = Frame(ML, MB, W - ML - MR, H - MT - MB,              id="later")

    doc = BaseDocTemplate(
        path, pagesize=A4,
        pageTemplates=[
            PageTemplate(id="First", frames=[first_frame], onPage=_draw_title),
            PageTemplate(id="Later", frames=[later_frame], onPage=_draw_header_footer),
        ],
    )

    # ── Styles ─────────────────────────────────────────────────────────────────
    base = ParagraphStyle(
        "RDBase", fontName="Helvetica", fontSize=11,
        leading=17, spaceAfter=6, alignment=TA_JUSTIFY,
    )
    hc = dict(alignment=TA_LEFT, textColor=accent, fontName="Helvetica-Bold")
    styles = {
        "h1":   ParagraphStyle("H1", parent=base, fontSize=18, spaceAfter=8,  spaceBefore=12, **hc),
        "h2":   ParagraphStyle("H2", parent=base, fontSize=14, spaceAfter=6,  spaceBefore=10, **hc),
        "h3":   ParagraphStyle("H3", parent=base, fontSize=12, spaceAfter=5,  spaceBefore=8,  **hc),
        "h4":   ParagraphStyle("H4", parent=base, fontSize=11, spaceAfter=4,  spaceBefore=6,  **hc),
        "para": base,
        "quote":    ParagraphStyle("Quote",    parent=base, leftIndent=20, rightIndent=20,
                                   fontName="Helvetica-Oblique", textColor=colors.HexColor("#555")),
        "bullet":   ParagraphStyle("Bullet",   parent=base, leftIndent=18, spaceAfter=3),
        "numbered": ParagraphStyle("Numbered", parent=base, leftIndent=18, spaceAfter=3),
        "code":     ParagraphStyle("Code", fontName="Courier", fontSize=9, leading=13,
                                   backColor=grey_bg, leftIndent=14, rightIndent=14,
                                   spaceAfter=8, spaceBefore=4, borderPad=5),
        "cell":     ParagraphStyle("Cell", parent=base, fontSize=9.5, spaceAfter=0),
    }

    # ── Build story ─────────────────────────────────────────────────────────────
    from reportlab.platypus import NextPageTemplate
    story: list = [NextPageTemplate("Later")]
    num_ctr: dict[int, int] = {}

    for blk in body_blocks:
        t = blk["type"]
        if t == "blank":
            story.append(Spacer(1, 4))
        elif t == "heading":
            story.append(Paragraph(_inline_pdf(blk["text"]), styles[f"h{min(blk['level'],4)}"]))
        elif t == "paragraph":
            story.append(Paragraph(_inline_pdf(blk["text"]), styles["para"]))
        elif t == "quote":
            story.append(Paragraph(_inline_pdf(blk["text"]), styles["quote"]))
        elif t == "bullet":
            story.append(Paragraph("• " + _inline_pdf(blk["text"]), styles["bullet"]))
        elif t == "numbered":
            d = blk.get("depth", 0)
            num_ctr[d] = num_ctr.get(d, 0) + 1
            for k in [k for k in num_ctr if k > d]: del num_ctr[k]
            story.append(Paragraph(f"{num_ctr[d]}. " + _inline_pdf(blk["text"]), styles["numbered"]))
        elif t == "code":
            story.append(Spacer(1, 4))
            story.append(Preformatted(_rl_sanitize("\n".join(blk["lines"])), styles["code"]))
        elif t == "hr":
            story.append(Spacer(1, 6))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#C8D8EA")))
            story.append(Spacer(1, 6))
        elif t == "image":
            img_bytes = _fetch_image_bytes(blk["src"])
            if img_bytes:
                from io import BytesIO
                from reportlab.platypus import Image as _RLImage
                try:
                    align     = _resolve_align(blk.get("align", "auto"), img_bytes)
                    w_inches  = _smart_width_inches(align, text_w=W - ML - MR)
                    w_pts     = w_inches  # already in points-compatible units (inch)
                    rl_img    = _RLImage(BytesIO(img_bytes), width=w_inches * inch, kind="bound")
                    rl_img.hAlign = {"left": "LEFT", "right": "RIGHT"}.get(align, "CENTER")
                    story.append(Spacer(1, 6))
                    story.append(rl_img)
                    if blk.get("alt"):
                        story.append(Paragraph(
                            _rl_sanitize(blk["alt"]),
                            ParagraphStyle("ImgCap", parent=base, fontSize=9,
                                           textColor=colors.HexColor("#777777"),
                                           alignment=TA_CENTER, spaceAfter=4),
                        ))
                    story.append(Spacer(1, 6))
                except Exception:
                    pass  # skip silently if image can't be embedded

        elif t == "table":
            rows = blk["rows"]
            if not rows: continue
            ncols = max(len(r) for r in rows)
            data  = []
            for ri, row in enumerate(rows):
                data.append([
                    Paragraph(
                        f"<b>{_escape_pdf(row[ci] if ci < len(row) else '')}</b>" if ri == 0
                        else _inline_pdf(row[ci] if ci < len(row) else ""),
                        styles["cell"],
                    )
                    for ci in range(ncols)
                ])
            tbl = Table(data, hAlign="LEFT", repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND",     (0, 0), (-1,  0), accent),
                ("TEXTCOLOR",      (0, 0), (-1,  0), colors.white),
                ("GRID",           (0, 0), (-1, -1), 0.5, colors.HexColor("#C8D8EA")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, grey_alt]),
                ("VALIGN",         (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING",     (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING",  (0, 0), (-1, -1), 5),
                ("LEFTPADDING",    (0, 0), (-1, -1), 6),
                ("RIGHTPADDING",   (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 6))
            story.append(tbl)
            story.append(Spacer(1, 6))

    doc.build(story)


# ══════════════════════════════════════════════════════════════════════════════
# HTML renderer
# ══════════════════════════════════════════════════════════════════════════════

def markdown_to_html(content: str, path: str) -> None:
    """Render Markdown to a self-contained, responsive HTML file."""
    title = Path(path).stem
    body  = _blocks_to_html_str(_parse_blocks(content))

    Path(path).write_text(
        f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
{_HTML_CSS}
</style>
</head>
<body>
<article>
{body}
</article>
</body>
</html>""",
        encoding="utf-8",
    )


_HTML_CSS = """
/* RAGdoll generated document */
*, *::before, *::after { box-sizing: border-box; }

:root {
  --accent:   #2E74B5;
  --accent2:  #1A4F8A;
  --bg:       #ffffff;
  --text:     #1a1a1a;
  --muted:    #555555;
  --code-bg:  #f4f4f4;
  --border:   #dde1e7;
  --th-bg:    #2E74B5;
  --th-text:  #ffffff;
  --tr-alt:   #f8fafd;
}

@media (prefers-color-scheme: dark) {
  :root {
    --bg:      #0d1117;
    --text:    #e6edf3;
    --muted:   #8b949e;
    --code-bg: #161b22;
    --border:  #30363d;
    --th-bg:   #1A4F8A;
    --tr-alt:  #161b22;
  }
}

body {
  font-family: "Segoe UI", system-ui, -apple-system, sans-serif;
  font-size: 16px;
  line-height: 1.7;
  color: var(--text);
  background: var(--bg);
  margin: 0;
  padding: 0;
}

article {
  max-width: 780px;
  margin: 2.5rem auto;
  padding: 0 1.5rem 4rem;
}

h1, h2, h3, h4, h5, h6 {
  color: var(--accent);
  font-weight: 700;
  line-height: 1.3;
  margin-top: 1.8em;
  margin-bottom: 0.5em;
}
h1 { font-size: 2.0em;  border-bottom: 2px solid var(--border); padding-bottom: .35em; }
h2 { font-size: 1.55em; border-bottom: 1px solid var(--border); padding-bottom: .25em; }
h3 { font-size: 1.25em; }
h4 { font-size: 1.1em; }

p  { margin: 0.75em 0; }

a           { color: var(--accent); text-decoration: underline; }
a:hover     { color: var(--accent2); }
strong      { font-weight: 700; }
em          { font-style: italic; }

code {
  font-family: "Cascadia Code", "Fira Code", Consolas, monospace;
  font-size: .875em;
  background: var(--code-bg);
  border: 1px solid var(--border);
  border-radius: 4px;
  padding: .1em .4em;
}

pre {
  background: var(--code-bg);
  border: 1px solid var(--border);
  border-radius: 8px;
  padding: 1rem 1.25rem;
  overflow-x: auto;
  margin: 1.25em 0;
}
pre code { background: none; border: none; padding: 0; font-size: .875em; line-height: 1.6; }

blockquote {
  border-left: 4px solid var(--accent);
  margin: 1.25em 0;
  padding: .5em 1em;
  color: var(--muted);
  font-style: italic;
}
blockquote p { margin: 0; }

ul, ol { padding-left: 1.6em; margin: .75em 0; }
li     { margin: .3em 0; }

hr {
  border: none;
  border-top: 1px solid var(--border);
  margin: 2em 0;
}

table {
  width: 100%;
  border-collapse: collapse;
  margin: 1.5em 0;
  font-size: .95em;
}
th, td { border: 1px solid var(--border); padding: .6em .9em; text-align: left; }
th     { background: var(--th-bg); color: var(--th-text); font-weight: 600; }
tr:nth-child(even) td { background: var(--tr-alt); }

figure { margin: 0; }
figure img { height: auto; border-radius: 6px; box-shadow: 0 2px 8px rgba(0,0,0,0.10); display: block; }
figcaption { font-size: 0.82em; color: var(--muted); margin-top: 0.4em; font-style: italic; text-align: center; }

figure.img-full  { width: 100%; margin: 1.5em 0; }
figure.img-full img { max-width: 100%; margin: 0 auto; }

figure.img-center { width: 68%; margin: 1.5em auto; }
figure.img-center img { width: 100%; }

figure.img-wrap-left {
  float: left;
  width: 42%;
  margin: 0.4em 1.6em 1em 0;
  clear: left;
}
figure.img-wrap-left img { width: 100%; }

figure.img-wrap-right {
  float: right;
  width: 42%;
  margin: 0.4em 0 1em 1.6em;
  clear: right;
}
figure.img-wrap-right img { width: 100%; }

h1, h2, h3, h4, h5, h6, hr { clear: both; }
"""


# ══════════════════════════════════════════════════════════════════════════════
# Dispatcher
# ══════════════════════════════════════════════════════════════════════════════

#: Extensions handled by this module (all others fall back to plain-text write)
FORMATTED_EXTS: frozenset[str] = frozenset({"docx", "doc", "pdf", "html", "htm"})


def render_document(ext: str, content: str, path: str) -> None:
    """Route Markdown content to the correct format renderer."""
    ext = ext.lower()
    if ext in ("docx", "doc"):
        markdown_to_docx(content, path)
    elif ext == "pdf":
        markdown_to_pdf(content, path)
    elif ext in ("html", "htm"):
        markdown_to_html(content, path)
    else:
        raise ValueError(f"No rich renderer for '.{ext}'. Supported: docx, pdf, html.")
